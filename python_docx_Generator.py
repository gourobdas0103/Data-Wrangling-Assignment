from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title of the Document
doc.add_heading('Data Wrangling Assignment Report', 0)

# Add Introduction
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "This report explains the workflow of two main components of the assignment: "
    "'TMA_data.sql' and 'Data_Wrangling.ipynb'. The 'TMA_data.sql' script creates and populates the database, "
    "while the 'Data_Wrangling.ipynb' Jupyter notebook performs the data wrangling, table creation, and visualization tasks."
)

# Section for TMA_data.sql
doc.add_heading('1. TMA_data.sql', level=1)

doc.add_paragraph(
    "The 'TMA_data.sql' script creates the database named 'TMA_data', defines the structure of the 'TMA_data' table, "
    "and populates the table with recruitment data. "
    "This table holds details about offers made and accepted for different locations and departments."
)

# Add SQL Code Description
doc.add_heading('SQL Code Explanation', level=2)
doc.add_paragraph(
    "The script first checks if the 'TMA_data' database exists, and if not, it creates the database. "
    "Then, the 'TMA_data' table is created with fields such as 'Location', 'Department', 'headcount', "
    "'Offers_Recruitment_Firm1', and other relevant columns. The recruitment data is inserted into this table. "
)

# Sample SQL Code Block
doc.add_paragraph(
    """
CREATE DATABASE IF NOT EXISTS TMA_data;
USE TMA_data;

CREATE TABLE IF NOT EXISTS TMA_data (
    Location VARCHAR(50),
    Department VARCHAR(50),
    headcount INT,
    Offers_Recruitment_Firm1 VARCHAR(50),
    Offers_Recruitment_Firm2 VARCHAR(50),
    Offers_Recruitment_Firm3 VARCHAR(50),
    Offers_Total VARCHAR(50),
    Acceptance_Recruitment_Firm1 VARCHAR(50),
    Acceptance_Recruitment_Firm2 VARCHAR(50),
    Acceptance_Recruitment_Firm3 VARCHAR(50),
    Acceptance_Total VARCHAR(50)
);

INSERT INTO `TMA_data`
    ...
    """
)

# Section for Data_Wrangling.ipynb
doc.add_heading('2. Data_Wrangling.ipynb', level=1)
doc.add_paragraph(
    "The 'Data_Wrangling.ipynb' notebook performs data wrangling, table creation, and visualization using MySQL and Python. "
    "It connects to the 'TMA_data' database, creates two new tables ('easy_data' and 'fig1'), fetches the data, "
    "and then visualizes the data using Matplotlib."
)

# Add Jupyter Notebook Code Explanation
doc.add_heading('Notebook Code Explanation', level=2)

# Add Connection Description
doc.add_paragraph(
    "1. MySQL Connection: A function is created to connect to the MySQL database using the MySQL Connector in Python."
)

doc.add_paragraph(
    """
def connect_to_db():
    connection = mysql.connector.connect(
        host='localhost',
        user='root',
        password='Toor',
        database='TMA_data'
    )
    return connection
    """
)

# Add easy_data Table Description
doc.add_paragraph(
    "2. Creating the 'easy_data' Table: A new table named 'easy_data' is created based on the 'TMA_data' table. "
    "This table contains key columns for easier analysis."
)

doc.add_paragraph(
    """
def create_easy_data_table():
    cursor.execute('''
    CREATE TABLE easy_data AS
    SELECT Location, Department, headcount AS Total_Headcount, Offers_Recruitment_Firm1, ...
    ''')
    """
)

# Add fig1 Table Description
doc.add_paragraph(
    "3. Creating the 'fig1' Table: Another table 'fig1' is created to aggregate the recruitment data for analysis. "
    "It sums up the offers made and accepted by location and department."
)

doc.add_paragraph(
    """
def create_fig1_table():
    cursor.execute('''
    CREATE TABLE fig1 AS
    SELECT Location, Department, SUM(headcount), ...
    ''')
    """
)

# Add Visualization Section
doc.add_heading('Data Visualization', level=2)
doc.add_paragraph(
    "The notebook fetches the data from the 'easy_data' table, and the data is displayed in a custom table format "
    "using Matplotlib. The table is visually styled to make it easy to interpret."
)

doc.add_paragraph(
    """
def display_custom_table(df):
    fig, ax = plt.subplots(figsize=(16, 8))
    ax.table(cellText=df.values, colLabels=df.columns, cellLoc='center', loc='center')
    plt.show()
    """
)

# Conclusion
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This assignment demonstrates the full workflow from setting up the database and populating it with data "
    "to performing data wrangling, creating new tables, and visualizing the results. "
    "The 'TMA_data.sql' file handles the database creation, while the 'Data_Wrangling.ipynb' "
    "provides the data processing and visualization pipeline."
)

# Save the document
doc.save('Data_Wrangling_Assignment_Report.docx')
